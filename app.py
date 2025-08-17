# pip install wxPython openpyxl pandas google-generativeai matplotlib seaborn python-docx reportlab

import wx
import wx.grid
import wx.dataview
from wx.lib.scrolledpanel import ScrolledPanel
import pandas as pd
import google.generativeai as genai
from datetime import datetime
import json
import os
import threading
import shutil

# --- Charting & Document Libraries ---
import matplotlib
matplotlib.use('WXAgg')
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as ReportLabImage, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet

class LoadingOverlay(wx.Panel):
	"""A semi-transparent overlay panel with a loading spinner and text."""
	def __init__(self, parent):
		super().__init__(parent)
		self.SetBackgroundColour(wx.Colour(255, 255, 255, 128))
		self.loading_text = wx.StaticText(self, label="Generating Report...\nCommunicating with Gemini AI.")
		self.spinner = wx.ActivityIndicator(self)
		sizer = wx.BoxSizer(wx.VERTICAL)
		sizer.AddStretchSpacer()
		sizer.Add(self.loading_text, 0, wx.CENTER | wx.BOTTOM, 15)
		sizer.Add(self.spinner, 0, wx.CENTER)
		sizer.AddStretchSpacer()
		self.SetSizer(sizer)
		self.Bind(wx.EVT_SIZE, self.on_size)
		self.Hide()

	def on_size(self, event):
		self.SetSize(self.GetParent().GetSize())
		event.Skip()

	def Show(self):
		super().Show()
		self.spinner.Start()

	def Hide(self):
		self.spinner.Stop()
		super().Hide()

class MainFrame(wx.Frame):
	def __init__(self):
		super().__init__(None, title="0zen", size=(1300, 900))

		icon = wx.Icon("logo.png", wx.BITMAP_TYPE_PNG)
		self.SetIcon(icon)

		self.excel_file = None
		self.df = None
		self.api_key = ""
		self.report_content = []
		self.chart_images = []
		self.wrappable_text_widgets = []

		self.panel = wx.Panel(self)
		self.splitter = wx.SplitterWindow(self.panel, style=wx.SP_LIVE_UPDATE)

		self.config_panel = wx.Panel(self.splitter, style=wx.BORDER_SUNKEN)
		self.output_panel = wx.Panel(self.splitter, style=wx.BORDER_SUNKEN)

		self.create_config_widgets(self.config_panel)
		self.create_output_widgets(self.output_panel)
		
		self.loading_overlay = LoadingOverlay(self.output_panel)

		self.splitter.SplitVertically(self.config_panel, self.output_panel, 550)
		self.splitter.SetMinimumPaneSize(400)

		main_sizer = wx.BoxSizer(wx.VERTICAL)
		main_sizer.Add(self.splitter, 1, wx.EXPAND)
		self.panel.SetSizer(main_sizer)

		self.SetMinSize((1000, 700))
		self.Centre()
		self.Show()
		self.Bind(wx.EVT_CLOSE, self.on_close)
		self.update_workflow_state()

	def create_config_widgets(self, parent):
		main_sizer = wx.BoxSizer(wx.VERTICAL)
		parent.SetSizer(main_sizer)
		title_font = wx.Font(16, wx.DEFAULT, wx.NORMAL, wx.BOLD)
		title_label = wx.StaticText(parent, label="Excel Report Generator")
		title_label.SetFont(title_font)
		main_sizer.Add(title_label, 0, wx.ALL | wx.CENTER, 15)
		self.status_bar = wx.StaticText(parent, label="Step 1: Please set your API key to begin.")
		self.status_bar.SetForegroundColour(wx.Colour(100, 100, 100))
		main_sizer.Add(self.status_bar, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
		api_sizer = self.create_api_section(parent)
		file_sizer = self.create_file_section(parent)
		report_sizer = self.create_report_config_section(parent)
		main_sizer.Add(api_sizer, 0, wx.EXPAND | wx.ALL, 10)
		main_sizer.Add(file_sizer, 0, wx.EXPAND | wx.ALL, 10)
		main_sizer.Add(report_sizer, 1, wx.EXPAND | wx.ALL, 10)
		
		
		main_sizer.AddStretchSpacer()
		dev_label = wx.StaticText(parent, label="Developed by 01one")
		font = dev_label.GetFont()
		font.SetPointSize(8)
		dev_label.SetFont(font)
		dev_label.SetForegroundColour(wx.Colour(128, 128, 128))
		main_sizer.Add(dev_label, 0, wx.ALIGN_RIGHT | wx.RIGHT | wx.BOTTOM, 5)		
		
		

	def create_api_section(self, parent):
		self.api_section_box = wx.StaticBox(parent, label="1. API Setup")
		sizer = wx.StaticBoxSizer(self.api_section_box, wx.VERTICAL)
		grid_sizer = wx.FlexGridSizer(2, 2, 5, 5)
		grid_sizer.AddGrowableCol(1)
		grid_sizer.Add(wx.StaticText(self.api_section_box, label="API Key:"), 0, wx.ALIGN_CENTER_VERTICAL)
		self.api_key_entry = wx.TextCtrl(self.api_section_box, style=wx.TE_PASSWORD)
		grid_sizer.Add(self.api_key_entry, 1, wx.EXPAND)
		grid_sizer.Add(wx.StaticText(self.api_section_box, label="Status:"), 0, wx.ALIGN_CENTER_VERTICAL)
		status_sizer = wx.BoxSizer(wx.HORIZONTAL)
		self.api_status_icon = wx.StaticText(self.api_section_box, label="🔴")
		self.api_status_label = wx.StaticText(self.api_section_box, label="Not Configured")
		status_sizer.Add(self.api_status_icon, 0, wx.ALIGN_CENTER_VERTICAL | wx.RIGHT, 5)
		status_sizer.Add(self.api_status_label, 1, wx.ALIGN_CENTER_VERTICAL)
		grid_sizer.Add(status_sizer, 1, wx.EXPAND)
		sizer.Add(grid_sizer, 1, wx.EXPAND | wx.ALL, 5)
		self.set_api_btn = wx.Button(self.api_section_box, label="Set API Key")
		sizer.Add(self.set_api_btn, 0, wx.EXPAND | wx.ALL, 5)
		self.Bind(wx.EVT_BUTTON, self.on_set_api_key, self.set_api_btn)
		return sizer

	def create_file_section(self, parent):
		self.file_section_box = wx.StaticBox(parent, label="2. Load Data")
		sizer = wx.StaticBoxSizer(self.file_section_box, wx.VERTICAL)
		self.file_path_label = wx.StaticText(self.file_section_box, label="No file selected.")
		self.browse_btn = wx.Button(self.file_section_box, label="Browse for Excel File...")
		sizer.Add(self.file_path_label, 0, wx.EXPAND | wx.ALL, 5)
		sizer.Add(self.browse_btn, 0, wx.EXPAND | wx.ALL, 5)
		self.Bind(wx.EVT_BUTTON, self.on_browse_file, self.browse_btn)
		return sizer

	def create_report_config_section(self, parent):
		self.report_config_box = wx.StaticBox(parent, label="3. Configure & Generate Report")
		sizer = wx.StaticBoxSizer(self.report_config_box, wx.VERTICAL)
		model_label = wx.StaticText(self.report_config_box, label="Select AI Model:")
		self.model_choice = wx.Choice(self.report_config_box, choices=[])
		sizer.Add(model_label, 0, wx.EXPAND | wx.TOP | wx.LEFT | wx.RIGHT, 5)
		sizer.Add(self.model_choice, 0, wx.EXPAND | wx.ALL, 5)
		report_type_label = wx.StaticText(self.report_config_box, label="Select a report type:")
		self.report_type_radio = wx.RadioBox(self.report_config_box, choices=["Sales Performance", "Market Analysis", "Employee Insights", "Custom Prompt"])
		sizer.Add(report_type_label, 0, wx.EXPAND | wx.TOP | wx.LEFT | wx.RIGHT, 5)
		sizer.Add(self.report_type_radio, 0, wx.EXPAND | wx.ALL, 5)
		self.custom_prompt_label = wx.StaticText(self.report_config_box, label="Your Custom Prompt:")
		self.custom_prompt_entry = wx.TextCtrl(self.report_config_box, style=wx.TE_MULTILINE, size=(-1, 80))
		self.custom_prompt_label.Hide()
		self.custom_prompt_entry.Hide()
		sizer.Add(self.custom_prompt_label, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.TOP, 5)
		sizer.Add(self.custom_prompt_entry, 1, wx.EXPAND | wx.ALL, 5)
		self.generate_btn = wx.Button(self.report_config_box, label="Generate Report")
		self.generate_btn.SetFont(wx.Font(10, wx.DEFAULT, wx.NORMAL, wx.BOLD))
		sizer.AddStretchSpacer()
		sizer.Add(self.generate_btn, 0, wx.EXPAND | wx.ALL, 5)
		self.Bind(wx.EVT_RADIOBOX, self.on_toggle_custom_prompt, self.report_type_radio)
		self.Bind(wx.EVT_BUTTON, self.on_generate_report, self.generate_btn)
		return sizer

	def create_output_widgets(self, parent):
		sizer = wx.BoxSizer(wx.VERTICAL)
		self.notebook = wx.Notebook(parent)

		# --- Data Preview Tab ---
		preview_panel = wx.Panel(self.notebook)
		preview_sizer = wx.BoxSizer(wx.VERTICAL)
		self.data_info_label = wx.StaticText(preview_panel, label="Load an Excel file to see a preview.")
		self.data_preview_ctrl = wx.dataview.DataViewListCtrl(preview_panel)
		preview_sizer.Add(self.data_info_label, 0, wx.ALL | wx.EXPAND, 5)
		preview_sizer.Add(self.data_preview_ctrl, 1, wx.EXPAND)
		preview_panel.SetSizer(preview_sizer)

		# --- Report Generation Tab ---
		report_panel = wx.Panel(self.notebook)
		report_sizer = wx.BoxSizer(wx.VERTICAL)
		button_sizer = wx.BoxSizer(wx.HORIZONTAL)
		self.save_docx_btn = wx.Button(report_panel, label="Save as Word (.docx)")
		self.save_pdf_btn = wx.Button(report_panel, label="Save as PDF (.pdf)")
		self.clear_btn = wx.Button(report_panel, label="Clear Output")
		button_sizer.Add(self.clear_btn, 0, wx.RIGHT, 5)
		button_sizer.Add(self.save_pdf_btn, 0, wx.RIGHT, 5)
		button_sizer.Add(self.save_docx_btn, 0)
		report_sizer.Add(button_sizer, 0, wx.ALL | wx.ALIGN_RIGHT, 5)
		self.report_display_panel = ScrolledPanel(report_panel)
		self.report_display_panel.SetupScrolling()
		self.report_display_sizer = wx.BoxSizer(wx.VERTICAL)
		self.report_display_panel.SetSizer(self.report_display_sizer)
		report_sizer.Add(self.report_display_panel, 1, wx.EXPAND)
		report_panel.SetSizer(report_sizer)

		self.notebook.AddPage(preview_panel, "📄 Data Preview")
		self.notebook.AddPage(report_panel, "📊 Generated Report")
		sizer.Add(self.notebook, 1, wx.EXPAND)
		parent.SetSizer(sizer)

		self.Bind(wx.EVT_BUTTON, lambda evt: self.on_save_report('docx'), self.save_docx_btn)
		self.Bind(wx.EVT_BUTTON, lambda evt: self.on_save_report('pdf'), self.save_pdf_btn)
		self.Bind(wx.EVT_BUTTON, self.on_clear_output, self.clear_btn)
		self.report_display_panel.Bind(wx.EVT_SIZE, self.on_report_panel_size)

	def on_report_panel_size(self, event):
		width = self.report_display_panel.GetClientSize().width - 20
		for widget in self.wrappable_text_widgets:
			widget.Wrap(width)
		self.report_display_sizer.Layout()
		event.Skip()

	def update_workflow_state(self):
		is_api_set = bool(self.api_key)
		is_data_loaded = self.df is not None
		self.file_section_box.Enable(is_api_set)
		self.report_config_box.Enable(is_data_loaded)
		if not is_api_set:
			self.status_bar.SetLabel("Step 1: Please set your API key to begin.")
		elif not is_data_loaded:
			self.status_bar.SetLabel("Step 2: Please load an Excel file.")
		else:
			self.status_bar.SetLabel("Step 3: Ready to generate your report.")
		self.config_panel.Layout()

	def on_set_api_key(self, event):
		api_key_value = self.api_key_entry.GetValue().strip()
		if not api_key_value:
			wx.MessageBox("API Key cannot be empty.", "Error", wx.OK | wx.ICON_ERROR); return
		self.api_section_box.Enable(False)
		self.api_status_label.SetLabel("Validating...")
		self.api_status_icon.SetLabel("⚪")
		threading.Thread(target=self._validate_api_key_thread, args=(api_key_value,), daemon=True).start()

	def _validate_api_key_thread(self, api_key_value):
		try:
			genai.configure(api_key=api_key_value)
			models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
			model_names = [name.split('/')[-1] for name in models]
			if not model_names:
				raise Exception("No compatible models found.")
			result = {'success': True, 'api_key': api_key_value, 'models': model_names}
		except Exception as e:
			result = {'success': False, 'error': e}
		wx.CallAfter(self._on_api_validation_complete, result)

	def _on_api_validation_complete(self, result):
		self.api_section_box.Enable(True)
		if result['success']:
			self.api_key = result['api_key']
			self.model_choice.SetItems(result['models'])
			default = "gemini-1.5-flash"
			if default in result['models']:
				self.model_choice.SetStringSelection(default)
			else:
				self.model_choice.SetSelection(0)
			self.api_status_icon.SetLabel("🟢")
			self.api_status_label.SetLabel("API Key Valid")
			wx.MessageBox("API Key configured!", "Success", wx.OK | wx.ICON_INFORMATION)
		else:
			self.api_key = ""
			self.api_status_icon.SetLabel("🔴")
			self.api_status_label.SetLabel("Invalid or Failed")
			self.model_choice.Clear()
			wx.MessageBox(f"Failed to connect.\n\nError: {result['error']}", "API Error", wx.OK | wx.ICON_ERROR)
		self.update_workflow_state()

	def on_browse_file(self, event):
		with wx.FileDialog(self, "Open Excel file", wildcard="Excel files (*.xlsx;*.xls)|*.xlsx;*.xls", style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST) as fileDialog:
			if fileDialog.ShowModal() == wx.ID_CANCEL: return
			self.excel_file = fileDialog.GetPath()
			self.file_path_label.SetLabel(f"Loaded: {os.path.basename(self.excel_file)}")
			self.load_excel_data()

	def load_excel_data(self):
		try:
			self.df = pd.read_excel(self.excel_file)
			self.data_info_label.SetLabel(f"Loaded {len(self.df)} rows, {len(self.df.columns)} columns from {os.path.basename(self.excel_file)}")
			self.update_data_preview()
			self.notebook.SetSelection(0)
		except Exception as e:
			self.df = None
			wx.MessageBox(f"Failed to load Excel file: {e}", "Error", wx.OK | wx.ICON_ERROR)
		finally:
			self.update_workflow_state()

	def update_data_preview(self):
		self.data_preview_ctrl.ClearColumns()
		self.data_preview_ctrl.DeleteAllItems()
		if self.df is None: return
		for col in self.df.columns: self.data_preview_ctrl.AppendTextColumn(col, width=120)
		for _, row in self.df.head(100).iterrows(): self.data_preview_ctrl.AppendItem([str(v) for v in row])

	def on_toggle_custom_prompt(self, event):
		is_custom = self.report_type_radio.GetStringSelection() == "Custom Prompt"
		self.custom_prompt_label.Show(is_custom)
		self.custom_prompt_entry.Show(is_custom)
		self.config_panel.Layout()

	def on_generate_report(self, event):
		if self.report_type_radio.GetStringSelection() == "Custom Prompt" and not self.custom_prompt_entry.GetValue().strip():
			wx.MessageBox("Please enter your custom prompt.", "Error"); return
		self.notebook.SetSelection(1)
		self.config_panel.Disable()
		self.loading_overlay.Show()
		report_type = self.report_type_radio.GetStringSelection()
		self.status_bar.SetLabel(f"Generating '{report_type}' report...")
		threading.Thread(target=self._generate_report_thread, daemon=True).start()

	def _generate_report_thread(self):
		try:
			wx.CallAfter(self.on_clear_output, confirm=False)
			data_summary = self.prepare_data_summary()
			prompt = self.create_prompt(data_summary)
			model_name = self.model_choice.GetStringSelection()
			model = genai.GenerativeModel(f"models/{model_name}")
			response = model.generate_content(prompt)
			cleaned_response = response.text.strip().replace('```json', '').replace('```', '')
			report_data = json.loads(cleaned_response)
			wx.CallAfter(self.process_and_display_report, report_data)
		except json.JSONDecodeError:
			raw_response = getattr(response, 'text', 'No response text available.')
			wx.CallAfter(self.display_error, f"AI returned an invalid JSON format. Raw response:\n\n{raw_response}")
		except Exception as e:
			wx.CallAfter(self.display_error, f"An error occurred: {e}")
		finally:
			wx.CallAfter(self.hide_loading_screen)

	def hide_loading_screen(self):
		self.loading_overlay.Hide()
		self.config_panel.Enable()
		self.update_workflow_state()

	def process_and_display_report(self, report_data):
		if 'title' in report_data:
			title = report_data['title']
			self.report_content.append({'type': 'h1', 'content': title})
			self.add_report_widget(wx.StaticText(self.report_display_panel, label=title), font=wx.Font(16, wx.DEFAULT, wx.NORMAL, wx.BOLD))
		for section in report_data.get('sections', []):
			sec_type = section.get('type')
			title = section.get('title', '')
			if title:
				self.report_content.append({'type': 'h2', 'content': title})
				self.add_report_widget(wx.StaticText(self.report_display_panel, label=title), font=wx.Font(12, wx.DEFAULT, wx.NORMAL, wx.BOLD), space_before=15)
			if sec_type == 'text':
				content = section.get('content', '')
				self.report_content.append({'type': 'text', 'content': content})
				text_widget = wx.StaticText(self.report_display_panel, label=content)
				self.wrappable_text_widgets.append(text_widget)
				self.add_report_widget(text_widget)
			elif sec_type == 'table':
				data = section.get('data', [])
				self.report_content.append({'type': 'table', 'data': data})
				if data: self.render_table_in_ui(data)
			elif sec_type == 'chart':
				try:
					chart_paths = self.create_chart(section)
					if chart_paths:
						self.report_content.append({'type': 'chart', 'paths': chart_paths, 'title': title})
						self.render_chart_in_ui(chart_paths)
						desc = section.get('description', '')
						if desc:
							self.report_content.append({'type': 'text', 'content': desc})
							desc_widget = wx.StaticText(self.report_display_panel, label=f"Chart Description: {desc}")
							desc_widget.SetForegroundColour(wx.Colour(100, 100, 100))
							self.wrappable_text_widgets.append(desc_widget)
							self.add_report_widget(desc_widget)
				except Exception as e:
					error_msg = f"Failed to generate chart '{title}': {e}"
					self.report_content.append({'type': 'error', 'content': error_msg})
					self.add_report_widget(wx.StaticText(self.report_display_panel, label=error_msg), color=wx.RED)
		self.on_report_panel_size(wx.SizeEvent(self.report_display_panel.GetSize()))
		self.report_display_panel.SetSizer(self.report_display_sizer)
		self.report_display_panel.SetupScrolling(scroll_x=False)

	def add_report_widget(self, widget, font=None, color=None, space_before=5, space_after=5):
		if font: widget.SetFont(font)
		if color: widget.SetForegroundColour(color)
		self.report_display_sizer.AddSpacer(space_before)
		self.report_display_sizer.Add(widget, 0, wx.EXPAND | wx.LEFT | wx.RIGHT, 10)
		self.report_display_sizer.AddSpacer(space_after)

	def render_table_in_ui(self, data):
		if not data or not isinstance(data, list) or len(data) < 1: return
		headers = data[0]
		rows = data[1:]
		num_rows = len(rows)
		num_cols = len(headers)
		grid = wx.grid.Grid(self.report_display_panel)
		grid.CreateGrid(num_rows, num_cols)
		for i, header in enumerate(headers):
			grid.SetColLabelValue(i, str(header))
		for r_idx, row_data in enumerate(rows):
			for c_idx, cell_data in enumerate(row_data):
				grid.SetCellValue(r_idx, c_idx, str(cell_data))
		grid.AutoSizeColumns()
		grid.EnableEditing(False)
		row_height = grid.GetDefaultRowSize()
		header_height = grid.GetColLabelSize()
		total_height = min((row_height * num_rows) + header_height + 5, 350)
		grid.SetMinSize(wx.Size(-1, total_height))
		self.report_display_sizer.Add(grid, 0, wx.EXPAND | wx.ALL, 10)
		self.report_display_sizer.Layout()

	def render_chart_in_ui(self, chart_paths):
		png_path = chart_paths.get('png')
		if not png_path or not os.path.exists(png_path): return
		bitmap = wx.Bitmap(png_path, wx.BITMAP_TYPE_PNG)
		img = bitmap.ConvertToImage()
		w, h = img.GetWidth(), img.GetHeight()
		max_width = self.report_display_panel.GetClientSize().width - 40
		if w > max_width and max_width > 0:
			new_h = int(h * (max_width / w))
			img = img.Scale(max_width, new_h, wx.IMAGE_QUALITY_HIGH)
			bitmap = img.ConvertToBitmap()
		img_widget = wx.StaticBitmap(self.report_display_panel, bitmap=bitmap)
		self.report_display_sizer.Add(img_widget, 0, wx.CENTER | wx.ALL, 10)

	def create_chart(self, chart_info):
		plt.figure(figsize=(8, 4.5))
		sns.set_theme(style="whitegrid")
		chart_type = chart_info.get('chart_type')
		x = chart_info.get('x_axis')
		if chart_type == 'hist':
			if not x: raise ValueError("Histogram requires an 'x_axis'.")
			if x not in self.df.columns: raise ValueError(f"Column for histogram not found in data: '{x}'.")
			sns.histplot(data=self.df, x=x)
			plt.ylabel("Frequency")
		else:
			y = chart_info.get('y_axis')
			if not all([chart_type, x, y]): raise ValueError("Chart info missing 'chart_type', 'x_axis', or 'y_axis'.")
			if x not in self.df.columns or y not in self.df.columns:
				raise ValueError(f"Columns for chart not found in data: '{x}' or '{y}'.")
			if chart_type == 'bar': sns.barplot(x=self.df[x], y=self.df[y])
			elif chart_type == 'line': sns.lineplot(x=self.df[x], y=self.df[y])
			elif chart_type == 'scatter': sns.scatterplot(x=self.df[x], y=self.df[y])
			else: raise ValueError(f"Unsupported chart type: {chart_type}")
			plt.ylabel(y)
		plt.title(chart_info.get('title', 'Chart'))
		plt.xlabel(x)
		plt.xticks(rotation=45, ha='right'); plt.tight_layout()
		timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
		if not os.path.exists('temp_charts'):
			os.makedirs('temp_charts')
		png_path = os.path.join('temp_charts', f"chart_{timestamp}.png")
		plt.savefig(png_path)
		plt.close()
		paths = {'png': png_path}
		self.chart_images.append(paths)
		return paths

	def prepare_data_summary(self):
		head_json = self.df.head(5).to_json(orient='records')
		with pd.option_context('display.max_columns', None):
			desc_json = self.df.describe(include='all').to_json()
		return json.dumps({
			"columns": list(self.df.columns),
			"data_types": {c: str(self.df[c].dtype) for c in self.df.columns},
			"data_head": json.loads(head_json),
			"description": json.loads(desc_json)
		}, indent=2)

	def create_prompt(self, data_summary):
		report_selection = self.report_type_radio.GetStringSelection()
		custom_prompt = self.custom_prompt_entry.GetValue()
		task = custom_prompt if report_selection == "Custom Prompt" else report_selection
		return f"""
		Analyze the provided Excel data summary and generate a professional business report.
		Your response MUST be a single valid JSON object. Do not include any text or markdown formatting like ```json before or after the JSON.
		The JSON root object must have a "title" (string) and "sections" (list of objects).
		
		**CRITICAL INSTRUCTIONS:**
		1.  **Multiple Viewpoints:** Analyze the data from multiple perspectives. If asked for a sales report, show sales by region, sales over time, and a table of top performers.
		2.  **Generate Various Charts:** Be proactive. If the data supports it, generate MULTIPLE charts. Include 'bar', 'line', and 'hist' charts if they are relevant.
		3.  **Insightful Analysis:** Each chart or table should be accompanied by a 'text' section or a 'description' that explains what the data shows and provides actionable insights.

		Each object in the "sections" list can be one of three types: 'text', 'table', or 'chart'.
		- For a 'text' section: "type": "text", "title": "Section Title", "content": "Paragraph text."
		- For a 'table' section: "type": "table", "title": "Table Title", "data": a 2D array with headers.
		- For a 'chart' section: "type": "chart", "title": "Chart Title", "description": "Brief analysis.", "chart_type": "bar", "x_axis": "column_for_x", "y_axis": "column_for_y".
		
		Supported chart types: 'bar', 'line', 'scatter', 'hist'.
		**IMPORTANT**: For a 'hist' chart, ONLY provide the 'x_axis'. The y-axis is automatically the frequency.
		
		Ensure column names for 'x_axis' and 'y_axis' exist in the data summary.
		
		Data Summary:
		{data_summary}
		
		---Report Task---
		{task}
		"""

	def display_error(self, error_message):
		self.on_clear_output(confirm=False)
		self.add_report_widget(wx.StaticText(self.report_display_panel, label="An Error Occurred"), font=wx.Font(12, wx.DEFAULT, wx.NORMAL, wx.BOLD), color=wx.RED)
		err_text = wx.StaticText(self.report_display_panel, label=error_message)
		self.wrappable_text_widgets.append(err_text)
		self.add_report_widget(err_text)
		self.report_display_panel.Layout()

	def on_save_report(self, file_format):
		if not self.report_content:
			wx.MessageBox("There is no report content to save.", "Warning", wx.OK | wx.ICON_WARNING); return
		initial_file = f"{self.report_type_radio.GetStringSelection().replace(' ', '_')}_report_{datetime.now().strftime('%Y%m%d')}"
		with wx.FileDialog(self, f"Save Report as {file_format.upper()}", wildcard=f"{file_format.upper()} files (*.{file_format})|*.{file_format}", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT, defaultFile=initial_file) as fileDialog:
			if fileDialog.ShowModal() == wx.ID_CANCEL: return
			path = fileDialog.GetPath()
			try:
				if file_format == 'docx': self.save_as_word(path)
				elif file_format == 'pdf': self.save_as_pdf(path)
				wx.MessageBox(f"Report saved to:\n{path}", "Success", wx.OK | wx.ICON_INFORMATION)
			except Exception as e:
				wx.MessageBox(f"Failed to save report: {e}", "Error", wx.OK | wx.ICON_ERROR)

	def save_as_word(self, file_path):
		doc = Document()
		for item in self.report_content:
			item_type = item.get('type')
			content = item.get('content', '')
			if item_type == 'h1':
				doc.add_heading(content, level=1)
			elif item_type == 'h2':
				doc.add_heading(content, level=2)
			elif item_type == 'text':
				doc.add_paragraph(content)
			elif item_type == 'chart' and item.get('paths', {}).get('png'):
				try:
					doc.add_picture(item['paths']['png'], width=Inches(6.0))
				except FileNotFoundError:
					doc.add_paragraph(f"[Chart not found at {item['paths']['png']}]")
			elif item_type == 'table' and item.get('data'):
				data = item['data']
				if not (isinstance(data, list) and len(data) > 0):
					continue
				
				headers = data
				records = data[1:]
				
				table = doc.add_table(rows=1, cols=len(headers))
				table.style = 'Table Grid'
				
				hdr_cells = table.rows[0].cells
				
				for i, header_text in enumerate(headers):
					hdr_cells[i].text = str(header_text)
					
				for record_data in records:
					row_cells = table.add_row().cells
					for i, cell_text in enumerate(record_data):
						if i < len(headers):
							row_cells[i].text = str(cell_text)
						
		doc.save(file_path)

	def save_as_pdf(self, file_path):
		doc = SimpleDocTemplate(file_path, pagesize=letter)
		styles = getSampleStyleSheet()
		story = []
		for item in self.report_content:
			item_type = item.get('type')
			content = item.get('content', '')
			if item_type == 'h1':
				story.append(Paragraph(content, styles['h1']))
			elif item_type == 'h2':
				story.append(Paragraph(content, styles['h2']))
			elif item_type == 'text':
				story.append(Paragraph(content, styles['BodyText']))
			elif item_type == 'chart' and item.get('paths'):
				png_path = item['paths']['png']
				if os.path.exists(png_path):
					img = ReportLabImage(png_path, width=6*inch, height=3.375*inch, hAlign='CENTER')
					story.append(img)
			elif item_type == 'table' and item.get('data'):
				t = Table(item['data'])
				t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey), ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
				story.append(t)
			story.append(Spacer(1, 12))
		doc.build(story)

	def on_clear_output(self, confirm=True):
		if confirm:
			if wx.MessageDialog(self, "Are you sure you want to clear the report output?", "Confirm Clear", wx.YES_NO | wx.ICON_QUESTION).ShowModal() != wx.ID_YES:
				return
		for widget in self.report_display_panel.GetChildren():
			widget.Destroy()
		self.report_content.clear()
		self.wrappable_text_widgets.clear()
		temp_dir = 'temp_charts'
		if os.path.exists(temp_dir):
			shutil.rmtree(temp_dir)
		self.chart_images.clear()
		self.report_display_sizer.Layout()

	def on_close(self, event):
		self.on_clear_output(confirm=False)
		self.Destroy()

if __name__ == "__main__":
	app = wx.App(False)
	frame = MainFrame()
	app.MainLoop()